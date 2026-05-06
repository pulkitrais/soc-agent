"""
Report Exporter
Generates investigation reports in Word, Excel, and HTML formats.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

import pandas as pd

from app.config import get_settings
from app.investigation.session import InvestigationSession

logger = logging.getLogger(__name__)


class ReportExporter:
    """Export investigation sessions to various report formats."""

    def __init__(self) -> None:
        self._settings = get_settings()
        self._exports_dir = self._settings.exports_dir
        self._exports_dir.mkdir(parents=True, exist_ok=True)

    # ── HTML ──────────────────────────────────────────────────────────────────

    def export_html(self, session: InvestigationSession) -> Path:
        """Export session to a self-contained dark-mode HTML report."""
        html = self._render_html(session)
        filename = f"investigation_{session.id[:8]}_{_safe_timestamp()}.html"
        path = self._exports_dir / filename
        path.write_text(html, encoding="utf-8")
        logger.info("HTML report written to %s", path)
        return path

    # ── Word ──────────────────────────────────────────────────────────────────

    def export_word(self, session: InvestigationSession) -> Path:
        """Export session to a Word (.docx) incident report."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise RuntimeError("python-docx is required for Word export.") from exc

        doc = Document()

        # Title
        title = doc.add_heading(f"Incident Report: {session.title}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        meta = doc.add_table(rows=4, cols=2)
        meta.style = "Table Grid"
        _set_table_row(meta, 0, "Session ID", session.id)
        _set_table_row(meta, 1, "Analyst", session.analyst)
        _set_table_row(meta, 2, "Severity", session.severity.upper())
        _set_table_row(meta, 3, "Status", session.status.upper())
        doc.add_paragraph()

        # Description
        if session.description:
            doc.add_heading("Description", level=1)
            doc.add_paragraph(session.description)

        # Notes
        if session.notes:
            doc.add_heading("Investigation Notes", level=1)
            doc.add_paragraph(session.notes)

        # Entities
        if session.entities:
            doc.add_heading("Observed Entities", level=1)
            for etype, values in session.entities.items():
                doc.add_paragraph(f"• {etype}: {', '.join(values)}")

        # Evidence
        doc.add_heading("Evidence", level=1)
        for i, ev in enumerate(session.evidence, start=1):
            doc.add_heading(f"{i}. {ev.title}", level=2)
            doc.add_paragraph(f"Type: {ev.item_type} | Time: {ev.timestamp}")
            if ev.content:
                p = doc.add_paragraph()
                run = p.add_run(ev.content[:2000])  # Truncate very long content
                run.font.name = "Courier New"
                run.font.size = Pt(8)

        # Timeline
        if session.timeline:
            doc.add_heading("Timeline", level=1)
            for event in session.timeline:
                doc.add_paragraph(
                    f"[{event.get('timestamp', '')}] {event.get('event', '')}",
                    style="List Bullet",
                )

        filename = f"investigation_{session.id[:8]}_{_safe_timestamp()}.docx"
        path = self._exports_dir / filename
        doc.save(str(path))
        logger.info("Word report written to %s", path)
        return path

    # ── Excel ─────────────────────────────────────────────────────────────────

    def export_excel(self, session: InvestigationSession) -> Path:
        """Export session data to Excel with multiple worksheets."""
        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("openpyxl is required for Excel export.") from exc

        filename = f"investigation_{session.id[:8]}_{_safe_timestamp()}.xlsx"
        path = self._exports_dir / filename

        with pd.ExcelWriter(str(path), engine="openpyxl") as writer:
            # Summary sheet
            summary_df = pd.DataFrame([
                {"Field": "Session ID", "Value": session.id},
                {"Field": "Title", "Value": session.title},
                {"Field": "Analyst", "Value": session.analyst},
                {"Field": "Severity", "Value": session.severity},
                {"Field": "Status", "Value": session.status},
                {"Field": "Created", "Value": session.created_at},
                {"Field": "Updated", "Value": session.updated_at},
                {"Field": "Description", "Value": session.description},
            ])
            summary_df.to_excel(writer, sheet_name="Summary", index=False)

            # Evidence sheet
            if session.evidence:
                evidence_df = pd.DataFrame([
                    {
                        "ID": ev.id[:8],
                        "Type": ev.item_type,
                        "Title": ev.title,
                        "Timestamp": ev.timestamp,
                        "Tags": ", ".join(ev.tags),
                        "Content Preview": ev.content[:500],
                    }
                    for ev in session.evidence
                ])
                evidence_df.to_excel(writer, sheet_name="Evidence", index=False)

            # Timeline sheet
            if session.timeline:
                timeline_df = pd.DataFrame(session.timeline)
                timeline_df.to_excel(writer, sheet_name="Timeline", index=False)

            # Entities sheet
            if session.entities:
                rows = []
                for etype, values in session.entities.items():
                    for v in values:
                        rows.append({"Type": etype, "Value": v})
                entities_df = pd.DataFrame(rows)
                entities_df.to_excel(writer, sheet_name="Entities", index=False)

            # Query results (preview only)
            qr_items = [e for e in session.evidence if e.item_type == "query_result"]
            if qr_items:
                for i, item in enumerate(qr_items[:10]):  # Limit to 10 query sheets
                    preview = item.data.get("preview", [])
                    if preview:
                        df = pd.DataFrame(preview)
                        sheet_name = f"Query_{i + 1}"[:31]  # Excel sheet name limit
                        df.to_excel(writer, sheet_name=sheet_name, index=False)

        logger.info("Excel report written to %s", path)
        return path

    # ── HTML template ─────────────────────────────────────────────────────────

    def _render_html(self, session: InvestigationSession) -> str:
        """Render a dark-mode HTML investigation report."""
        severity_colors = {
            "critical": "#ff4444",
            "high": "#ff8800",
            "medium": "#ffcc00",
            "low": "#44bb44",
        }
        sev_color = severity_colors.get(session.severity, "#888888")

        evidence_html = ""
        for i, ev in enumerate(session.evidence, start=1):
            content_escaped = (
                ev.content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            evidence_html += f"""
            <div class="evidence-item">
                <h3>{i}. {ev.title}</h3>
                <span class="badge">{ev.item_type}</span>
                <span class="meta">{ev.timestamp}</span>
                <pre><code>{content_escaped[:3000]}</code></pre>
            </div>"""

        entities_html = ""
        for etype, values in session.entities.items():
            entities_html += f"<li><strong>{etype}:</strong> {', '.join(values)}</li>"

        timeline_html = "".join(
            f'<li><code>{e.get("timestamp", "")}</code> — {e.get("event", "")}</li>'
            for e in session.timeline
        )

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Sentinel Investigator Report — {session.title}</title>
  <style>
    :root {{ --bg: #0d1117; --surface: #161b22; --border: #30363d; --text: #e6edf3; --accent: #58a6ff; }}
    body {{ background: var(--bg); color: var(--text); font-family: 'Segoe UI', sans-serif; padding: 2rem; }}
    h1 {{ color: var(--accent); border-bottom: 2px solid var(--border); padding-bottom: .5rem; }}
    h2 {{ color: #8b949e; margin-top: 2rem; }}
    h3 {{ color: var(--accent); }}
    .meta-table {{ border-collapse: collapse; margin: 1rem 0; }}
    .meta-table td {{ padding: .4rem .8rem; border: 1px solid var(--border); }}
    .badge {{ background: #21262d; border: 1px solid var(--border); border-radius: 1rem;
              padding: .1rem .6rem; font-size: .8rem; margin-right: .5rem; }}
    .severity {{ color: {sev_color}; font-weight: bold; }}
    .evidence-item {{ background: var(--surface); border: 1px solid var(--border);
                       border-radius: .5rem; padding: 1rem; margin: 1rem 0; }}
    pre {{ background: #010409; border-radius: .3rem; padding: 1rem; overflow-x: auto; }}
    code {{ font-family: 'Cascadia Code', monospace; font-size: .85rem; color: #79c0ff; }}
    ul {{ line-height: 1.8; }}
    .meta {{ color: #8b949e; font-size: .8rem; }}
  </style>
</head>
<body>
  <h1>🔍 Sentinel Investigator — Incident Report</h1>
  <table class="meta-table">
    <tr><td>Title</td><td><strong>{session.title}</strong></td></tr>
    <tr><td>Session ID</td><td><code>{session.id}</code></td></tr>
    <tr><td>Analyst</td><td>{session.analyst}</td></tr>
    <tr><td>Severity</td><td class="severity">{session.severity.upper()}</td></tr>
    <tr><td>Status</td><td>{session.status}</td></tr>
    <tr><td>Created</td><td>{session.created_at}</td></tr>
    <tr><td>Updated</td><td>{session.updated_at}</td></tr>
  </table>

  {"<h2>Description</h2><p>" + session.description + "</p>" if session.description else ""}
  {"<h2>Notes</h2><pre>" + session.notes + "</pre>" if session.notes else ""}

  <h2>Observed Entities</h2>
  <ul>{entities_html or "<li>None recorded</li>"}</ul>

  <h2>Evidence ({len(session.evidence)} items)</h2>
  {evidence_html or "<p>No evidence recorded.</p>"}

  <h2>Timeline</h2>
  <ul>{timeline_html or "<li>No events recorded</li>"}</ul>

  <footer style="color:#8b949e;margin-top:3rem;font-size:.8rem;">
    Generated by Sentinel Investigator v{get_settings().app_version}
    on {datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")}
  </footer>
</body>
</html>"""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _safe_timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


def _set_table_row(table, row_idx: int, key: str, value: str) -> None:
    """Helper to fill a 2-column Word table row."""
    row = table.rows[row_idx]
    row.cells[0].text = key
    row.cells[1].text = str(value)
